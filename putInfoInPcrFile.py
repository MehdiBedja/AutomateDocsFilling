from docx import Document

# Load the input document
input_path = r"C:\Users\DELL\Downloads\PCR MODEL.docx"
output_path =  r"C:\Users\DELL\Downloads\PCR_MODEL_OUTPUT.docx"
doc = Document(input_path)

# Define the static information
nom = "MUHAMMAD AWAIS"
prenom = "ABUBAKAR"
date_de_naissance = "01/01/1980"
resultat = "NEGATIF"

# Find and fill the first table
table1 = doc.tables[0]
# Fill Full Name
table1.cell(0, 0).text = f"Full Name: {nom} {prenom}"
# Fill Date of Birth
table1.cell(0, 1).text = f"Date Of Birth: {date_de_naissance}"

# Find and fill the second table
table2 = doc.tables[1]
# Fill 'pcr' in the second line second column
table2.cell(1, 1).text = "pcr"
# Fill the result in the second line first column
table2.cell(1, 3).text = f"Result: {resultat}"

# Save the output document
doc.save(output_path)
