import os
from datetime import datetime
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

def read_exposant_list(file_path):
    # Read the table data from the Word document
    doc = Document(file_path)
    table = doc.tables[0]

    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = [cell.text.strip() for cell in row.cells]
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    
    return data

def create_person_document(template_path, output_folder, person_info):
    # Load the template document
    doc = Document(template_path)

    # Fill in the information in the first table
    table1 = doc.tables[0]
    # Full Name
    full_name = f"{person_info.get('Nom', '')} {person_info.get('PRENOM', '')}"
    table1.cell(0, 0).text += f" {full_name}"
    # Date of Birth
    date_de_naissance = person_info.get('DATE DE NAISSANCE', '')
    table1.cell(0, 1).text += f" {date_de_naissance}"

    # Address
    if 'Adresse' in person_info:
        addresse = person_info['Adresse']
        table1.cell(1, 1).text += f" {addresse}"

    # Passport Number
    if 'NUMERO PASSPORT' in person_info:
        passport_num = person_info['NUMERO PASSPORT']
        table1.cell(1, 0).text += f" {passport_num}"

    # Nationality
    if 'NATIONALITE' in person_info:
        nationalite = person_info['NATIONALITE']
        table1.cell(0, 2).text += f" {nationalite}"

    # Fill in the information in the second table
    table2 = doc.tables[1]
    # PCR type
    table2.cell(1, 1).text += " pcr"
    # Result
    resultat = person_info.get('RESULTAT', '')
    table2.cell(1, 3).text += f" {resultat}"

    current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    table2.cell(1, 4).text += f" {current_datetime}"

    # Define the output file path
    output_path = os.path.join(output_folder, f"{full_name}.docx")
    # Save the document
    doc.save(output_path)

def process_documents(exposant_list_path, template_path, output_folder):
    # Read the exposant list
    exposant_list = read_exposant_list(exposant_list_path)

    # Create a document for each person
    for person_info in exposant_list:
        create_person_document(template_path, output_folder, person_info)

def select_exposant_list():
    file_path = filedialog.askopenfilename(title="Select Exposant List", filetypes=[("Word Documents", "*.docx")])
    exposant_list_entry.delete(0, tk.END)
    exposant_list_entry.insert(0, file_path)

def select_template():
    file_path = filedialog.askopenfilename(title="Select Template", filetypes=[("Word Documents", "*.docx")])
    template_entry.delete(0, tk.END)
    template_entry.insert(0, file_path)

def select_output_folder():
    folder_path = filedialog.askdirectory(title="Select Output Folder")
    output_folder_entry.delete(0, tk.END)
    output_folder_entry.insert(0, folder_path)

def start_processing():
    exposant_list_path = exposant_list_entry.get()
    template_path = template_entry.get()
    output_folder = output_folder_entry.get()

    if not exposant_list_path or not template_path or not output_folder:
        messagebox.showwarning("Input Error", "Please select all required files and folders.")
        return

    process_documents(exposant_list_path, template_path, output_folder)
    messagebox.showinfo("Completed", "Document creation completed successfully.")
    
    # Open the output folder in the file explorer
    os.startfile(output_folder)

# Create the main window
root = tk.Tk()
root.title("Document Processor")

# Create and place the GUI components
tk.Label(root, text="Exposant List:").grid(row=0, column=0, padx=10, pady=10)
exposant_list_entry = tk.Entry(root, width=50)
exposant_list_entry.grid(row=0, column=1, padx=10, pady=10)
tk.Button(root, text="Browse...", command=select_exposant_list).grid(row=0, column=2, padx=10, pady=10)

tk.Label(root, text="Template:").grid(row=1, column=0, padx=10, pady=10)
template_entry = tk.Entry(root, width=50)
template_entry.grid(row=1, column=1, padx=10, pady=10)
tk.Button(root, text="Browse...", command=select_template).grid(row=1, column=2, padx=10, pady=10)

tk.Label(root, text="Output Folder:").grid(row=2, column=0, padx=10, pady=10)
output_folder_entry = tk.Entry(root, width=50)
output_folder_entry.grid(row=2, column=1, padx=10, pady=10)
tk.Button(root, text="Browse...", command=select_output_folder).grid(row=2, column=2, padx=10, pady=10)

tk.Button(root, text="Start Processing", command=start_processing).grid(row=3, column=0, columnspan=3, padx=10, pady=20)

# Run the main event loop
root.mainloop()
