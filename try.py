import os
from tkinter import Tk, Label, Button, filedialog
from docx import Document
from docx.shared import Pt
import pandas as pd

def read_exposant_list(file_path):
    doc = Document(file_path)
    table = doc.tables[0]

    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    
    return data

def create_fiche_pcr(template_path, output_folder, person):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        if 'Nom' in paragraph.text:
            paragraph.text = f"Nom: {person['Nom']}"
            paragraph.style.font.size = Pt(12)
        if 'Prénom' in paragraph.text:
            paragraph.text = f"Prénom: {person['PRENOM']}"
            paragraph.style.font.size = Pt(12)
        if 'Date de naissance' in paragraph.text:
            paragraph.text = f"Date de naissance: {person['DATE DE NAISSANCE']}"
            paragraph.style.font.size = Pt(12)
        if 'Résultat' in paragraph.text:
            paragraph.text = f"Résultat: {person['RESULTAT']}"
            paragraph.style.font.size = Pt(12)

    output_path = os.path.join(output_folder, f"{person['Nom']}_{person['PRENOM']}.docx")
    doc.save(output_path)

def main():
    root = Tk()
    root.title("Fiche PCR Generator")
    root.geometry("300x200")

    def select_exposant_list():
        file_path = filedialog.askopenfilename(
            title="Select the Exposant List File",
            filetypes=[("Word Documents", "*.docx")])
        exposant_list_label.config(text=os.path.basename(file_path))
        root.exposant_list_path = file_path

    def select_output_folder():
        folder_path = filedialog.askdirectory(title="Select the Output Folder")
        output_folder_label.config(text=os.path.basename(folder_path))
        root.output_folder = folder_path

    def generate_files():
        fiche_pcr_template_path = filedialog.askopenfilename(
            title="Select the Fiche PCR Template",
            filetypes=[("Word Documents", "*.docx")])
        
        if hasattr(root, 'exposant_list_path') and hasattr(root, 'output_folder'):
            exposant_list = read_exposant_list(root.exposant_list_path)
            for person in exposant_list:
                create_fiche_pcr(fiche_pcr_template_path, root.output_folder, person)
            done_label.config(text="Files Generated Successfully")
        else:
            done_label.config(text="Please select all required files and folders")

    exposant_list_label = Label(root, text="Select Exposant List File")
    exposant_list_label.pack(pady=5)
    select_exposant_list_button = Button(root, text="Browse", command=select_exposant_list)
    select_exposant_list_button.pack(pady=5)

    output_folder_label = Label(root, text="Select Output Folder")
    output_folder_label.pack(pady=5)
    select_output_folder_button = Button(root, text="Browse", command=select_output_folder)
    select_output_folder_button.pack(pady=5)

    generate_button = Button(root, text="Generate Files", command=generate_files)
    generate_button.pack(pady=20)

    done_label = Label(root, text="")
    done_label.pack(pady=5)

    root.mainloop()

if __name__ == '__main__':
    main()
