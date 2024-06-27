from datetime import datetime
from docx import Document
import os

def read_exposant_list(file_path):
    # Read the table data from the Word document
    doc = Document(file_path)
    table = doc.tables[0]

    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = [cell.text.strip() for cell in row.cells]
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    
    return data



def create_person_document(template_path, output_folder, person_info):
    # Load the template document
    doc = Document(template_path)

    # Fill in the information in the first table
    table1 = doc.tables[0]
    # Full Name
    full_name = f"{person_info.get('Nom', '')} {person_info.get('PRENOM', '')}"
    table1.cell(0, 0).text += f" {full_name}"
    # Date of Birth
    date_de_naissance = person_info.get('DATE DE NAISSANCE', '')
    table1.cell(0, 1).text += f" {date_de_naissance}"

    # Address
    if 'ADRESSE' in person_info:
        addresse = person_info['ADRESSE']
        table1.cell(1, 1).text += f" {addresse}"

    # Passport Number
    if 'NUMERO PASSPORT' in person_info:
        passport_num = person_info['NUMERO PASSPORT']
        table1.cell(1, 0).text += f" {passport_num}"

    # Nationality
    if 'NATIONALITE' in person_info:
        nationalite = person_info['NATIONALITE']
        table1.cell(0, 2).text += f" {nationalite}"

    # Fill in the information in the second table
    table2 = doc.tables[1]
    # PCR type
    table2.cell(1, 1).text += " pcr"
    # Result
    resultat = person_info.get('RESULTAT', '')
    table2.cell(1, 3).text += f" {resultat}"

    current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    table2.cell(1, 4).text += f" {current_datetime}"

    # Define the output file path
    output_path = os.path.join(output_folder, f"{full_name}.docx")
    # Save the document
    doc.save(output_path)

def main():
    exposant_list_path = r"C:\Users\DELL\Downloads\EXPOSANT_LISTE_5.docx"  # Input exposant list file path
    template_path = r"C:\Users\DELL\Downloads\PCR MODEL.docx"              # Template file path
    output_folder = r"C:\Users\DELL\Downloads\output_docs"                 # Output folder path

    # Ensure the output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Read the exposant list
    exposant_list = read_exposant_list(exposant_list_path)
    #print(exposant_list)

    # Create a document for each person
    for person_info in exposant_list:
        create_person_document(template_path, output_folder, person_info)

if __name__ == '__main__':
    main()
